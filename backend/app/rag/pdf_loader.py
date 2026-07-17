from pypdf import PdfReader
from docx import Document


def load_pdf_text(file_path: str) -> str:
    reader = PdfReader(file_path)
    text = ""

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text.strip()


def load_docx_text(file_path: str) -> str:
    document = Document(file_path)
    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def load_text_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read().strip()


def load_document_text(file_path: str) -> str:
    lower_path = file_path.lower()

    if lower_path.endswith(".pdf"):
        return load_pdf_text(file_path)

    if lower_path.endswith(".docx"):
        return load_docx_text(file_path)

    if lower_path.endswith(".txt"):
        return load_text_file(file_path)

    return ""
